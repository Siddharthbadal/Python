#install python docx

from docx import Document
from docx.shared import Inches
document = Document()

#profile image
document.add_picture('img.png', width=Inches(2.0))

# title details
name = input("Enter your name: ")
phone = input("Enter your phone: ")
email = input("Enter your email: ")

document.add_paragraph(f"{name} | {phone} |{email}")


# about me
document.add_heading('About Me')
about_me = input('Tell about yourself: ')
document.add_paragraph(about_me)

# work exp.
document.add_heading('Work Expirence')
p = document.add_paragraph()

company = input("Enter company name:  ")
title = input(" Enter your job title:  ")
from_date = input('from: ')
to_date = input("To date: ")
p.add_run(company + ' ' +'\n').bold = True
p.add_run(title + ' ' +'\n').bold = True 
'\n'
p.add_run('From: '+ from_date + ' - ' + 'To: '+ to_date +'\n').italic = True
experience_details = input('Work responsibility at ' + company + ' ')
p.add_run(experience_details)

#more experiencs 
while True:
    more_experiences = input('Do you want to add one more company experience? Y or N: ')
    if more_experiences.lower() == 'y':
        p = document.add_paragraph()
        company = input("Enter company name:  ")
        title = input(" Enter your job title:  ")
        from_date = input('from: ')
        to_date = input("To date: ")
        p.add_run(company + ' ' +'\n').bold = True
        p.add_run(title + ' ').bold = True 
        '\n'
        p.add_run(from_date + ' - ' + to_date +'\n').italic = True
        experience_details = input('Work responsibility at ' + company + ' ')
        p.add_run(experience_details)
    else:
        break


# skils
document.add_heading('Skills')
skill = input('Enter your skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input('Do you want to add more skills? Y or n: ')
    if has_more_skills.lower() == 'y':
        skill = input('Enter your skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break 


document.add_heading('Education')
p = document.add_paragraph()

edu = input("Enter Recent Education details:  ")
college = input(" Enter your college:  ")
edu_from_date = input('from: ')
edu_to_date = input("To date: ")
p.add_run(edu + ' ' +'\n').bold = True
p.add_run(college + ' ' +'\n').bold = True 
'\n'
p.add_run('From: '+ edu_from_date + ' - ' + 'To: '+ edu_to_date +'\n').italic = True
edu_details = input('Course Details in ' + edu + ' ')
p.add_run(edu_details)

while True:
    more_edu = input('Do you want to add more education details? Y or N: ')
    if more_edu.lower() == 'y':
        p = document.add_paragraph()
        edu = input("Enter Recent Education details:  ")
        college = input(" Enter your college:  ")
        edu_from_date = input('from: ')
        edu_to_date = input("To date: ")
        p.add_run(edu + ' ' +'\n').bold = True
        p.add_run(college + ' ' +'\n').bold = True 
        '\n'
        p.add_run('From: '+ edu_from_date + ' - ' + ' To ' + edu_to_date +'\n').italic = True
        edu_details = input('Course Details in ' + edu + ' ')
        p.add_run(edu_details)
    else:
        break




#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "@sidbad using Python and DocX"


document.save('cv.docx')